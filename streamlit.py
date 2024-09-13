import os
from dotenv import load_dotenv
from openai import OpenAI
from transformers import BertTokenizer, BertForSequenceClassification
import torch
from docx import Document as WordDocument
import re
import streamlit as st
from audio_recorder_streamlit import audio_recorder
import tempfile
import random
import io

# Load environment variables from .env file
load_dotenv()

# Retrieve the API key from the environment
api_key = os.getenv("OPENAI_API_KEY")

# Initialize the OpenAI client with the API key
OpenAI.api_key = api_key
client = OpenAI()

# Load BERT model and tokenizer
tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
model = BertForSequenceClassification.from_pretrained('bert-base-uncased')

# Define a list of technical interview questions
interview_questions = [
    "Explain the concept of overfitting in machine learning.",
    "How does a decision tree algorithm work?",
    "What is the difference between supervised and unsupervised learning?",
    "Can you explain the bias-variance tradeoff?",
    "What is cross-validation and why is it used?"
]

def transcribe_audio(audio_file_path):
    try:
        with open(audio_file_path, 'rb') as audio_file:
            if not audio_file_path.lower().endswith(('.wav', '.mp3', '.mp4', '.m4a', '.mpeg', '.mpga', '.oga', '.ogg', '.webm')):
                raise ValueError("Unsupported file format. Please use one of the following formats: .wav, .mp3, .mp4, .m4a, .mpeg, .mpga, .oga, .ogg, .webm")

            transcription = client.audio.transcriptions.create(model="whisper-1", file=audio_file, response_format='text')
            return transcription
    except ValueError as e:
        print(f"Value error: {e}")
        return None
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return None

def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are a helpful assistant."},
                  {"role": "user", "content": "Summarize the following text in a concise abstract paragraph:\n" + transcription}],
        max_tokens=150
    )
    return response.choices[0].message.content.strip()

def key_points_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are a helpful Machine Learning Interviewing assistant."},
                  {"role": "user", "content": "Extract the main points from the text:\n" + transcription}],
        max_tokens=150
    )
    return response.choices[0].message.content.strip()

def evaluate_clarity_of_key_points(key_points):
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are a helpful Machine Learning Interviewing assistant."},
                  {"role": "user", "content": "For the following key points, rate the clarity on a scale from 1 to 5. Provide explanations for each rating:\n" + key_points}],
        max_tokens=150
    )
    return response.choices[0].message.content.strip()

def evaluate_relevance_of_key_points(key_points, question):
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are a helpful Machine Learning Interviewing assistant."},
                  {"role": "user", "content": "For the following key points, rate the relevance on a scale from 1 to 5. Provide explanations for each rating and compare the relevance with the interview question. Check if the points elaborate on the question or just repeat keywords:\n" + key_points + "\nInterview Question:\n" + question}],
        max_tokens=150
    )
    return response.choices[0].message.content.strip()

def evaluate_depth_of_key_points(key_points):
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are a helpful Machine Learning Interviewing assistant."},
                  {"role": "user", "content": "For the following key points, rate the depth of information on a scale from 1 to 5. Provide explanations for the rating and be stringent in the evaluation:\n" + key_points}],
        max_tokens=150
    )
    return response.choices[0].message.content.strip()

def evaluate_sentiment_of_transcription(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are a helpful sentiment analysis assistant."},
                  {"role": "user", "content": "Perform sentiment analysis on the following text. Rate the sentiment on a scale from 1 to 5 and provide explanations. Consider potential nervousness or filler words:\n" + transcription}],
        max_tokens=150
    )
    return response.choices[0].message.content.strip()

def split_complex_responses(response):
    tokens = tokenizer(response, return_tensors="pt", truncation=True, padding=True)
    outputs = model(**tokens)
    logits = outputs.logits
    probabilities = torch.nn.functional.softmax(logits, dim=1)
    _, predicted_class = torch.max(probabilities, dim=1)
    split_responses = re.split(r'(?<=\.)\s', response)
    return split_responses

def save_as_docx(minutes, filename):
    doc = WordDocument()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()
    doc.save(filename)

def process_audio_and_generate_report(audio_file, question):
    # Save the audio file to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_audio_file:
        temp_audio_file.write(audio_file.getvalue())
        temp_audio_file.flush()
        temp_audio_path = temp_audio_file.name

    transcription = transcribe_audio(temp_audio_path)
    if transcription is None:
        return None, None

    summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    clarity = evaluate_clarity_of_key_points(key_points)
    relevance = evaluate_relevance_of_key_points(key_points, question)
    depth = evaluate_depth_of_key_points(key_points)
    sentiment = evaluate_sentiment_of_transcription(transcription)

    parsed_responses = split_complex_responses(transcription)

    minutes = {
        'interview_question': question,
        'transcription': transcription,
        'abstract_summary': summary,
        'key_points': key_points,
        'clarity': clarity,
        'relevance': relevance,
        'depth': depth,
        'sentiment': sentiment,
        'parsed_responses': '\n'.join(parsed_responses)
    }

    filename = 'answer_evaluations.docx'
    save_as_docx(minutes, filename)
    return filename, minutes

def generate_interview_question():
    return random.choice(interview_questions)

st.title("Audio Processing and Report Generation")

if 'question' not in st.session_state:
    st.session_state.question = None

if 'minutes' not in st.session_state:
    st.session_state.minutes = None
    st.session_state.filename = None

def start_interview():
    st.session_state.question = generate_interview_question()

if st.button("Start Interview"):
    start_interview()
    st.write(f"**Interview Question:** {st.session_state.question}")

if st.session_state.question:
    action = st.radio("Choose an action", ["Record Audio", "Upload Audio"])

    if action == "Record Audio":
        st.header("Record Audio")
        st.write("Click the button below to start recording.")
        audio_bytes = audio_recorder(pause_threshold=2.0, sample_rate=41_000)
        
        if audio_bytes:
            question = st.session_state.question
            with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_audio_file:
                temp_audio_file.write(audio_bytes)
                temp_audio_file.flush()
                st.write("Processing recorded audio...")
                with open(temp_audio_file.name, 'rb') as audio_file:
                    filename, minutes = process_audio_and_generate_report(io.BytesIO(audio_bytes), question)
                    if filename:
                        st.session_state.filename = filename
                        st.session_state.minutes = minutes
                        st.write("**Transcription and Evaluation Results:**")
                        st.write(minutes)
                        st.download_button("Download DOCX", data=open(filename, 'rb').read(), file_name=filename, mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        st.success("Audio processed and document generated.")
                    else:
                        st.error("Failed to process audio.")

    elif action == "Upload Audio":
        st.header("Audio File Upload")
        audio_file = st.file_uploader("Upload an audio file", type=["wav", "mp3"])
        
        if audio_file is not None:
            st.audio(audio_file, format='audio/wav')
            
            if st.button("Process Audio"):
                question = st.session_state.question
                filename, minutes = process_audio_and_generate_report(io.BytesIO(audio_file.read()), question)
                if filename:
                    st.session_state.filename = filename
                    st.session_state.minutes = minutes
                    st.write("**Transcription and Evaluation Results:**")
                    st.write(minutes)
                    st.download_button("Download DOCX", data=open(filename, 'rb').read(), file_name=filename, mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    st.success("Audio processed and document generated.")
                else:
                    st.error("Failed to process audio.")